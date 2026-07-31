"""파일 → 평문 텍스트 추출 (필요하면 OCR).

PDF 는 먼저 텍스트 레이어를 시도하고, 거의 안 나오면 스캔본으로 보고 페이지를
이미지로 렌더링해 tesseract 에 넘긴다. 이미지 파일은 곧바로 OCR 한다.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

from ..config import DOCX_EXTS, IMAGE_EXTS, PDF_EXTS, TEXT_EXTS

log = logging.getLogger("chatchat.rag.extract")

# OCR 렌더링 배율. 300dpi 근처가 되도록 잡는다 (72dpi 기준 * 4)
OCR_SCALE = 4.0
OCR_MAX_PAGES = 50
OCR_MIN_SHORT_SIDE = 1400  # 이보다 짧은 변은 확대한다
OCR_PSM_LADDER = (6, 4, 11)  # 단일 블록 → 다단 → 성긴 텍스트
OCR_MIN_ACCEPT = 10  # 이 이상 나오면 그 PSM 결과를 채택

Progress = Callable[[int, int, str], None]  # (done, total, phase)


class Unsupported(Exception):
    """인덱싱 대상이 아닌 확장자."""


class OcrUnavailable(Exception):
    """tesseract 가 설치돼 있지 않거나 언어 데이터가 없음."""


@dataclass
class Extracted:
    text: str
    ocr: bool = False  # OCR 을 실제로 썼는지


def extract(path: Path, cfg: dict, on_progress: Progress | None = None) -> Extracted:
    ext = path.suffix.lower()
    if ext in TEXT_EXTS:
        return Extracted(_text(path))
    if ext in DOCX_EXTS:
        return Extracted(_docx(path))
    if ext in PDF_EXTS:
        return _pdf(path, cfg, on_progress)
    if ext in IMAGE_EXTS:
        if not cfg.get("rag_ocr_enabled"):
            raise Unsupported("이미지 OCR 이 꺼져 있습니다")
        return Extracted(_ocr_image(path, str(cfg["rag_ocr_langs"])), ocr=True)
    raise Unsupported(f"지원하지 않는 확장자입니다: {ext or '(없음)'}")


# ─────────────────────────── 평문 / docx ───────────────────────────


def _text(path: Path) -> str:
    raw = path.read_bytes()
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass
    # UTF-8 이 아니면 인코딩을 추정한다 (CP949 한글 문서가 흔하다)
    from charset_normalizer import from_bytes

    best = from_bytes(raw).best()
    return str(best) if best else raw.decode("utf-8", errors="replace")


def _docx(path: Path) -> str:
    import docx

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ─────────────────────────── PDF ───────────────────────────


def _pdf(path: Path, cfg: dict, on_progress: Progress | None) -> Extracted:
    text = _pdf_text_layer(path)

    min_chars = int(cfg.get("rag_ocr_min_chars", 80))
    if len(text.strip()) >= min_chars or not cfg.get("rag_ocr_enabled"):
        return Extracted(text)

    log.info("%s: 텍스트 레이어가 %d자뿐 → 스캔본으로 보고 OCR", path.name, len(text.strip()))
    try:
        return Extracted(_ocr_pdf(path, str(cfg["rag_ocr_langs"]), on_progress), ocr=True)
    except OcrUnavailable:
        raise
    except Exception as exc:  # noqa: BLE001 - OCR 실패 시 텍스트 레이어라도 살린다
        log.warning("PDF OCR 실패 %s: %s", path.name, exc)
        return Extracted(text)


def _pdf_text_layer(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001 - 페이지 하나가 깨져도 나머지는 살린다
            log.warning("pdf 페이지 추출 실패 %s p.%d: %s", path.name, i + 1, exc)
    return "\n\n".join(pages)


def _ocr_pdf(path: Path, langs: str, on_progress: Progress | None) -> str:
    import pypdfium2 as pdfium

    _ensure_tesseract()
    pdf = pdfium.PdfDocument(str(path))
    total = min(len(pdf), OCR_MAX_PAGES)
    if len(pdf) > OCR_MAX_PAGES:
        log.warning("%s: %d쪽 중 앞 %d쪽만 OCR 합니다", path.name, len(pdf), OCR_MAX_PAGES)

    out: list[str] = []
    try:
        for i in range(total):
            if on_progress:
                on_progress(i, total, "ocr")
            page = pdf[i]
            try:
                image = page.render(scale=OCR_SCALE).to_pil()
                out.append(_ocr_pil(image, langs))
            except Exception as exc:  # noqa: BLE001
                log.warning("OCR 실패 %s p.%d: %s", path.name, i + 1, exc)
            finally:
                page.close()
        if on_progress:
            on_progress(total, total, "ocr")
    finally:
        pdf.close()
    return "\n\n".join(out)


# ─────────────────────────── 이미지 OCR ───────────────────────────


def _ocr_image(path: Path, langs: str) -> str:
    from PIL import Image

    _ensure_tesseract()
    with Image.open(path) as image:
        return _ocr_pil(image, langs)


def _ocr_pil(image, langs: str) -> str:
    import pytesseract
    from PIL import Image, ImageOps

    if image.mode not in ("L", "RGB"):
        image = image.convert("RGB")
    image = ImageOps.exif_transpose(image) or image

    # 해상도가 낮으면 인식률이 급락한다. 짧은 변을 기준으로 키운다.
    short = min(image.size)
    if short < OCR_MIN_SHORT_SIDE:
        ratio = OCR_MIN_SHORT_SIDE / max(short, 1)
        image = image.resize(
            (int(image.width * ratio), int(image.height * ratio)), Image.LANCZOS
        )

    image = ImageOps.autocontrast(ImageOps.grayscale(image))

    # 기본 PSM 3(자동 분할)은 여백이 많은 페이지를 통째로 "Empty page" 로 버리는 일이
    # 잦다. 단일 블록(6) → 다단(4) → 성긴 텍스트(11) 순으로 내려가며 시도한다.
    for psm in OCR_PSM_LADDER:
        try:
            text = pytesseract.image_to_string(
                image, lang=langs, config=f"--psm {psm} --dpi 300"
            ).strip()
        except Exception as exc:  # noqa: BLE001 - 다음 PSM 으로 계속
            log.debug("OCR psm=%s 실패: %s", psm, exc)
            continue
        if len(text) >= OCR_MIN_ACCEPT:
            return text
    return ""


_tesseract_ok: bool | None = None


def _ensure_tesseract() -> None:
    global _tesseract_ok
    if _tesseract_ok:
        return
    import pytesseract

    try:
        pytesseract.get_tesseract_version()
        _tesseract_ok = True
    except Exception as exc:  # noqa: BLE001
        raise OcrUnavailable(f"tesseract 를 실행할 수 없습니다: {exc}") from exc


def available_langs() -> list[str]:
    try:
        import pytesseract

        return sorted(pytesseract.get_languages(config=""))
    except Exception:  # noqa: BLE001
        return []
